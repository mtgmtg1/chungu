#!/usr/bin/env python3
# [Flow: Step 1 (로그인) -> Step 2 (DOCX 샘플 생성) -> Step 3 (upload) -> Step 4 (confirm) -> Step 5 (poll) -> Step 6 (download)]
import base64
import http.client
import io
import json
import os
import tempfile
import time
import uuid
from pathlib import Path

import requests
from docx import Document

BASE_URL = os.environ.get("CHUNGU_BASE_URL", "192.168.1.50:28181")
SUPABASE_URL = os.environ.get("CHUNGU_SUPABASE_URL", "192.168.1.50:28000")
ANON_KEY = os.environ.get("CHUNGU_SUPABASE_ANON_KEY", "")
EMAIL = os.environ.get("CHUNGU_TEST_EMAIL", "test@chungu.local")
PASSWORD = os.environ.get("CHUNGU_TEST_PASSWORD", "Test1234!")


def _request(method: str, host: str, path: str, headers: dict = None, body: bytes = None, json_data=None):
    if json_data is not None:
        body = json.dumps(json_data).encode()
        headers = {**(headers or {}), "Content-Type": "application/json"}
    headers = headers or {}
    conn = http.client.HTTPConnection(host)
    conn.request(method, path, body=body, headers=headers)
    resp = conn.getresponse()
    data = resp.read()
    conn.close()
    text = data.decode("utf-8", errors="replace")
    if resp.status >= 400:
        raise RuntimeError(f"HTTP {resp.status}: {text[:500]}")
    return resp.status, text


def login(email: str, password: str) -> str:
    path = "/auth/v1/token?grant_type=password"
    _, text = _request(
        "POST",
        SUPABASE_URL,
        path,
        headers={"apikey": ANON_KEY},
        json_data={"email": email, "password": password},
    )
    return json.loads(text)["access_token"]


def _sample_docx() -> Path:
    doc = Document()
    doc.add_heading("Docling 백엔드 통합 테스트", level=1)
    doc.add_paragraph("이 DOCX 파일은 Docling 전처리 파이프라인을 검증합니다.")
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "항목"
    table.cell(0, 1).text = "값"
    table.cell(1, 0).text = "A"
    table.cell(1, 1).text = "100"
    table.cell(2, 0).text = "B"
    table.cell(2, 1).text = "200"
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        doc.save(tmp.name)
        return Path(tmp.name)


def upload_docx(token: str, sample_path: Path, refinement: bool = False):
    data = sample_path.read_bytes()
    boundary = uuid.uuid4().hex
    body = io.BytesIO()
    for name, value in {
        "pipeline": "vision",
        "columns": "항목,값",
        "prompt": "표를 정확히 유지해주세요",
        "dpi": "150",
        "docling_refinement": "true" if refinement else "false",
    }.items():
        body.write(f"--{boundary}\r\n".encode())
        body.write(f'Content-Disposition: form-data; name="{name}"\r\n\r\n'.encode())
        body.write(f"{value}\r\n".encode())
    body.write(f"--{boundary}\r\n".encode())
    body.write(b'Content-Disposition: form-data; name="files"; filename="test.docx"\r\n')
    body.write(b"Content-Type: application/vnd.openxmlformats-officedocument.wordprocessingml.document\r\n\r\n")
    body.write(data)
    body.write(b"\r\n")
    body.write(f"--{boundary}--\r\n".encode())
    payload = body.getvalue()

    headers = {
        "Authorization": f"Bearer {token}",
        "Content-Type": f"multipart/form-data; boundary={boundary}",
    }
    status, text = _request("POST", BASE_URL, "/api/jobs/upload", headers=headers, body=payload)
    print("upload status:", status)
    print("upload body:", text[:500])
    return json.loads(text)


def confirm_job(token: str, job_id: str):
    headers = {"Authorization": f"Bearer {token}"}
    status, text = _request("POST", BASE_URL, f"/api/jobs/{job_id}/confirm", headers=headers)
    print("confirm status:", status)
    print("confirm body:", text[:500])
    return json.loads(text)


def poll_job(token: str, job_id: str, max_seconds: int = 120):
    headers = {"Authorization": f"Bearer {token}"}
    for i in range(max_seconds):
        _, text = _request("GET", BASE_URL, f"/api/jobs/{job_id}", headers=headers)
        j = json.loads(text)
        print(f"[{i}s] status={j.get('status')} done_files={j.get('done_files')}/{j.get('total_files')}")
        if j.get("status") in ("completed", "done", "failed", "error"):
            return j
        time.sleep(2)
    return None


def download_result(job_id: str, token: str):
    headers = {"Authorization": f"Bearer {token}"}
    for ext in ("md", "csv"):
        status, text = _request("GET", BASE_URL, f"/api/jobs/{job_id}/download/{ext}", headers=headers)
        print(f"download {ext} status:", status)
        print(f"download {ext} preview:", text[:500])


def main():
    token = login(EMAIL, PASSWORD)
    print("token acquired")
    sample = _sample_docx()
    try:
        job = upload_docx(token, sample, refinement=True)
        job_id = job["job_id"]
        print("job_id:", job_id)
        confirm_job(token, job_id)
        result = poll_job(token, job_id)
        print("result:", json.dumps(result, default=str, indent=2)[:1000])
        download_result(job_id, token)
    finally:
        sample.unlink(missing_ok=True)


if __name__ == "__main__":
    main()
