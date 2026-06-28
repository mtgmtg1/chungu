#!/usr/bin/env python3
# [Flow: Step 1 (b2 Docling 서비스 health) -> Step 2 (DOCX 샘플 생성) -> Step 3 (/convert/file) -> Step 4 (이미지 다운로드) -> Step 5 (결과 출력)]
import json
import os
import tempfile
from pathlib import Path

import requests
from docx import Document

DOCLING_SERVICE_URL = os.environ.get("DOCLING_SERVICE_URL", "http://192.168.1.100:28182")


def _sample_docx() -> Path:
    doc = Document()
    doc.add_heading("Docling 테스트 문서", level=1)
    doc.add_paragraph("이 문서는 Docling 전처리 서비스의 동작을 확인하기 위한 샘플입니다.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "항목"
    table.cell(0, 1).text = "값"
    table.cell(1, 0).text = "A"
    table.cell(1, 1).text = "100"
    doc.add_paragraph("이미지 아래에 표가 있습니다.")
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        doc.save(tmp.name)
        return Path(tmp.name)


def test_health():
    url = f"{DOCLING_SERVICE_URL}/health"
    resp = requests.get(url, timeout=10)
    resp.raise_for_status()
    print("health:", resp.json())


def test_convert_file(sample_path: Path):
    url = f"{DOCLING_SERVICE_URL}/convert/file"
    with open(sample_path, "rb") as f:
        files = {"file": (sample_path.name, f, "application/octet-stream")}
        resp = requests.post(url, files=files, timeout=1200)
    resp.raise_for_status()
    data = resp.json()
    print("convert:", json.dumps({k: v for k, v in data.items() if k != "markdown"}, indent=2, ensure_ascii=False))
    print("markdown preview:\n", data.get("markdown", "")[:500])
    return data


def test_download_image(image_rel_path: str):
    url = f"{DOCLING_SERVICE_URL}/images/{image_rel_path}"
    resp = requests.get(url, timeout=60)
    resp.raise_for_status()
    print(f"image {image_rel_path}: {len(resp.content)} bytes")


def main():
    test_health()
    sample = _sample_docx()
    try:
        result = test_convert_file(sample)
        for rel in result.get("images", [])[:3]:
            test_download_image(rel)
    finally:
        sample.unlink(missing_ok=True)


if __name__ == "__main__":
    main()
